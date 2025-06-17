#!/usr/bin/env python3
import docx
import sys

def extract_docx_content(file_path):
    try:
        doc = docx.Document(file_path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        # Also extract table content if any
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))
        
        return "\n".join(content)
    except Exception as e:
        return f"Error reading document: {str(e)}"

if __name__ == "__main__":
    file_path = "/home/ubuntu/upload/PRDforAutoapplyappCSWPBC.docx"
    content = extract_docx_content(file_path)
    print(content)

